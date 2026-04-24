"""
JARVIS — Autonomous Assignment Report Generator
Produces a fully formatted Word document (.docx) for the Safe House
Network Design assignment (75 points, due 21 Apr 2026).
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = os.path.dirname(os.path.abspath(__file__))

IMG1 = os.path.join(OUT, "diagram1_physical_LAN.png")
IMG2 = os.path.join(OUT, "diagram2_logical_LAN.png")
IMG3 = os.path.join(OUT, "diagram3_logical_WAN.png")


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
    return p


def add_para(doc, text, bold=False, italic=False, fontsize=11, colour=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(fontsize)
    if colour:
        run.font.color.rgb = RGBColor(*colour)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_bullet(doc, text, fontsize=11):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(fontsize)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_img(doc, path, width_inches=6.3, caption=""):
    doc.add_picture(path, width=Inches(width_inches))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cp = doc.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.runs[0].italic = True
        cp.runs[0].font.size = Pt(9.5)
        cp.paragraph_format.space_after = Pt(14)


def add_table_row(table, cells, bold_first=False, bg=None):
    row = table.add_row()
    for i, (cell, text) in enumerate(zip(row.cells, cells)):
        run = cell.paragraphs[0].add_run(text)
        run.font.size = Pt(9.5)
        if bold_first and i == 0:
            run.bold = True
        if bg:
            set_cell_bg(cell, bg)
    return row


def build_report():
    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Cover block ───────────────────────────────────────────────────────────
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("Network Protocols and Architectures")
    tr.bold = True
    tr.font.size = Pt(22)
    tr.font.color.rgb = RGBColor(13, 71, 161)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Safe House Network Design — Assignment Report")
    sr.font.size = Pt(14)
    sr.font.color.rgb = RGBColor(84, 110, 122)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Module: Network Protocols and Architectures  |  Tutor: Dr. Seyedali Pourmoafi\n"
                 "Weighting: 75%  |  Word Count: ~2,200  |  Date: April 2026")
    meta.runs[0].font.size = Pt(10)
    doc.add_paragraph()

    # ── Introduction ──────────────────────────────────────────────────────────
    add_heading(doc, "Introduction", level=1, color=(13, 71, 161))
    add_para(doc,
        "This report presents a comprehensive network design for a high-security "
        "safe house operated by a UK-based intelligence and cyber-operations agency. "
        "The design addresses two interconnected tasks: a Local Area Network (LAN) "
        "for the safe house building itself, and a Wide Area Network (WAN) optical "
        "link connecting the safe house to the agency's main premises in London. "
        "All design decisions are driven by the core mandate of 100% availability "
        "under any single point of failure, enterprise-grade security, and scalability "
        "for future operational expansion."
    )

    # ═══════════════════════════════════════════════════════════════════════════
    #  CHAPTER 1 — LAN
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "Chapter 1: LAN Design", level=1, color=(13, 71, 161))

    # 1.1 Physical diagram
    add_heading(doc, "1.1 Physical LAN Diagram", level=2)
    add_para(doc,
        "Figure 1 illustrates the physical network layout overlaid on the safe house "
        "floor plan. The building contains twelve operational rooms arranged in a four-row, "
        "three-column grid. Room 8, designated the Communications Room (Comms Room), "
        "houses the complete core network and security infrastructure within a full-size "
        "42U rack. Room 5 serves as the primary panic room, connected to the physical "
        "security VLAN."
    )
    add_img(doc, IMG1, width_inches=6.2,
            caption="Figure 1 — Physical LAN Diagram: Safe House Floor Plan with Network Infrastructure")

    # 1.2 Logical diagram
    add_heading(doc, "1.2 Logical LAN Diagram", level=2)
    add_para(doc,
        "Figure 2 presents the logical three-tier hierarchical topology, showing VLAN "
        "segmentation, dual uplink paths, and the border security layer. The logical "
        "view abstracts physical room positions and focuses on data-flow relationships "
        "between network tiers."
    )
    add_img(doc, IMG2, width_inches=6.2,
            caption="Figure 2 — Logical LAN Diagram: Three-Layer Hierarchical Model with VLAN Segmentation")

    # 1.3 Design description
    add_heading(doc, "1.3 Design Description", level=2)

    add_heading(doc, "Three-Layer Hierarchical Model", level=3)
    add_para(doc,
        "The design adopts Cisco's three-layer hierarchical model (Oppenheimer, 1999), "
        "which separates network functions into Core, Distribution, and Access layers. "
        "This separation of concerns enables each layer to be optimised independently, "
        "simplifies troubleshooting, and provides a clear scalability path — additional "
        "rooms or devices can be added at the access layer without redesigning the core "
        "(Cisco Systems, 2022a)."
    )

    add_heading(doc, "Core Layer — Room 8 (Comms Room)", level=3)
    add_para(doc,
        "Two Cisco Catalyst 9500-48Y4C switches form a StackWise Virtual cluster in "
        "Room 8, appearing as a single logical switch to the network whilst remaining "
        "two physically separate units. This architecture eliminates the core switch as "
        "a single point of failure: if either physical unit experiences a hardware fault, "
        "the surviving unit continues to forward all traffic without reconfiguration and "
        "with sub-second failover (Cisco Systems, 2023a). The 25 Gbps virtual stack link "
        "between the two units ensures state synchronisation at all times."
    )
    add_para(doc,
        "The 4.8 Tbps aggregate switching capacity of the Catalyst 9500 provides "
        "substantial headroom beyond current requirements (48 workstations at 1 GbE "
        "equates to 96 Gbps aggregate), ensuring the core is never a bottleneck and "
        "accommodates future expansion without hardware replacement."
    )

    add_heading(doc, "Distribution Layer", level=3)
    add_para(doc,
        "Three Cisco Catalyst 9300-48P distribution switches aggregate the access layer, "
        "each serving a zone of four rooms (Dist-A: Rooms 1–4; Dist-B: Rooms 5–8; "
        "Dist-C: Rooms 9–12). Critically, each distribution switch maintains dual uplinks "
        "— one 10 GbE OS2 single-mode fibre path to Core-1, and a separate path to "
        "Core-2 — bonded via Equal-Cost Multi-Path (ECMP) routing. Should either core "
        "switch or either uplink fibre fail, traffic automatically redistributes to the "
        "surviving path with zero downtime (Cisco Systems, 2022b)."
    )
    add_para(doc,
        "The 9300-48P provides 860 W of PoE+ budget per switch, powering IP cameras, "
        "electronic door locks, and access points directly from the network — eliminating "
        "separate power cabling and the single points of failure it would introduce."
    )

    add_heading(doc, "Access Layer", level=3)
    add_para(doc,
        "Each of the twelve rooms contains a dedicated Cisco Catalyst 9200L-48P access "
        "switch providing four PoE+ ports for workstations, one port for the room's "
        "electronic door lock, and one port for any surveillance camera. The 9200L's "
        "managed switching capability supports full VLAN assignment per port, ensuring "
        "that workstation traffic (VLAN 10), physical security traffic (VLAN 40), and "
        "surveillance traffic (VLAN 20 or 30) are segregated at the point of connection. "
        "If an access switch in one room fails, only that room is affected — all other "
        "eleven rooms retain full network access."
    )

    add_heading(doc, "Redundancy — Meeting 100% Availability", level=3)
    add_para(doc,
        "The design satisfies the three explicitly stated failure scenarios:"
    )

    # Redundancy table
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    hdr[0].paragraphs[0].add_run("Failure Scenario").bold = True
    hdr[1].paragraphs[0].add_run("How the Design Responds").bold = True
    set_cell_bg(hdr[0], "0D47A1")
    set_cell_bg(hdr[1], "0D47A1")
    for cell in hdr:
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
    rows_data = [
        ("One link fails",
         "ECMP distributes load across both uplinks; if one is cut, "
         "the surviving uplink carries all traffic. RSTP (802.1w) reconverges in <1 second."),
        ("One switch fails",
         "StackWise Virtual core means loss of one core unit is invisible to the network. "
         "Distribution switches retain one surviving uplink. Only the affected access "
         "switch's room loses local connectivity."),
        ("One network path fails",
         "Pre-computed alternate paths via ECMP and VRRP maintain gateway reachability. "
         "VRRP failover completes in under 3 seconds, transparent to end users."),
    ]
    for i, (f, r) in enumerate(rows_data):
        bg = "E3F2FD" if i % 2 == 0 else "FFFFFF"
        add_table_row(tbl, [f, r], bold_first=True, bg=bg)
    doc.add_paragraph()

    add_heading(doc, "VLAN Segmentation", level=3)
    add_para(doc,
        "Traffic is separated into six VLANs to enforce security boundaries and ensure "
        "that compromise of one network segment cannot directly affect others:"
    )
    vlan_data = [
        ("VLAN 10", "Operations",  "All 48 workstations (12 rooms × 4 PCs)"),
        ("VLAN 20", "CCTV-A",      "Surveillance Circuit A — independent switch, isolated path"),
        ("VLAN 30", "CCTV-B",      "Surveillance Circuit B — independent switch, isolated path"),
        ("VLAN 40", "PhySec",      "Electronic locks, panic room control systems"),
        ("VLAN 50", "Management",  "Out-of-band switch and router management"),
        ("VLAN 60", "WAN/Border",  "Firewall ↔ border router ↔ external circuits"),
    ]
    vtbl = doc.add_table(rows=1, cols=3)
    vtbl.style = "Table Grid"
    vhdr = vtbl.rows[0].cells
    for ci, h in enumerate(["VLAN ID", "Name", "Purpose"]):
        vhdr[ci].paragraphs[0].add_run(h).bold = True
        set_cell_bg(vhdr[ci], "1B5E20")
        vhdr[ci].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    for i, (vid, vname, vdesc) in enumerate(vlan_data):
        bg = "E8F5E9" if i % 2 == 0 else "FFFFFF"
        add_table_row(vtbl, [vid, vname, vdesc], bold_first=True, bg=bg)
    doc.add_paragraph()

    add_heading(doc, "CCTV — Dual Independent Circuits", level=3)
    add_para(doc,
        "The two surveillance circuits are implemented as entirely independent network "
        "paths. Circuit A aggregates on a dedicated Cisco Catalyst 9200 switch in Room 8 "
        "on VLAN 20; Circuit B aggregates on a separate switch on VLAN 30. The two "
        "switches share only the core layer for uplink, on separate physical ports. "
        "An attacker who physically severs one circuit's switch or cabling cannot affect "
        "the other. Each circuit has its own UPS battery backup. This architecture "
        "satisfies the requirement that 'if one surveillance or network circuit is "
        "compromised, the remaining circuit must continue operating and be able to trigger "
        "alerts and backup response'."
    )

    add_heading(doc, "Physical Security — Electronic Locks and Lockdown", level=3)
    add_para(doc,
        "All electronic door locks and panic room control systems reside on VLAN 40, "
        "isolated from operational traffic. The lockdown controller in Room 8 connects "
        "via dual Ethernet paths to both core switches, ensuring lockdown commands reach "
        "all access switches even when one core unit has failed. In a breach event, the "
        "controller broadcasts a simultaneous lockdown command across VLAN 40 to all "
        "twelve access switches, engaging every electronic lock within milliseconds. "
        "The PoE+ connection also means that locks cannot be defeated by cutting a "
        "separate power cable — power and control share the same secured Cat6A runs "
        "inside steel conduit."
    )

    add_heading(doc, "Cabling", level=3)
    add_para(doc,
        "OS2 single-mode fibre (ITU-T G.652.D) with 10GBase-LR SFP+ transceivers forms "
        "the backbone between Room 8 and each distribution zone, supporting 10 Gbps at "
        "distances up to 10 km — far exceeding the building's 30 m dimensions and "
        "providing complete future-proofing. Cat6A UTP (TIA-568-C.2) connects access "
        "switches to workstations and peripherals, supporting 10 GbE at distances up to "
        "100 m. All cable runs are enclosed in steel conduit, consistent with the "
        "building's steel-reinforced physical security posture."
    )

    add_heading(doc, "Scalability", level=3)
    add_para(doc,
        "The hierarchical design scales in three dimensions without architectural "
        "redesign: (1) Additional rooms are served by adding access switches and "
        "connecting them to the nearest distribution switch — unused ports already "
        "exist. (2) Additional bandwidth is provided by upgrading the core-to-distribution "
        "links from 10 GbE to 25 GbE using the existing SFP+ module slots. "
        "(3) Additional external links (WAN or ISP) are added as new interfaces on the "
        "border router, with no changes to the LAN topology."
    )

    # 1.4 Hardware
    add_heading(doc, "1.4 Network Hardware — Critical Evaluation", level=2)

    hw_items = [
        ("Cisco Catalyst 9500-48Y4C — Core Switches",
         "The Catalyst 9500 is Cisco's purpose-built fixed-configuration core switch "
         "for enterprise and government environments. StackWise Virtual clustering "
         "delivers hitless failover in under 300 ms, making it the industry-standard "
         "choice when 100% availability is mandated (Cisco Systems, 2023a). The "
         "Network Advantage licence includes full Layer 3 routing (OSPF, BGP, MPLS), "
         "enabling seamless integration with the WAN layer. Dual field-replaceable "
         "power supplies (N+1) and hot-swappable fans ensure that no hardware "
         "maintenance window causes downtime. Alternatives such as the Juniper EX9200 "
         "offer comparable redundancy, but the Cisco ecosystem provides tighter "
         "integration with the 9300 and 9200L distribution and access tiers."),

        ("Cisco Catalyst 9300-48P — Distribution Switches",
         "The 9300 Series provides 48 PoE+ ports with an 860 W budget, eliminating "
         "the need for separate power infrastructure for cameras and electronic locks "
         "(Cisco Systems, 2022b). The modular network module accepts four 10 GbE SFP+ "
         "uplinks, which are cabled as dual uplinks to both core switches. Cisco's "
         "StackPower technology allows multiple 9300 units to share a common power "
         "pool, providing additional power resilience in the distribution zone."),

        ("Cisco Catalyst 9200L-48P — Access Switches",
         "The 9200L provides cost-effective managed access switching with PoE+ support "
         "and full VLAN/QoS capabilities. Its compact 1RU form factor fits within "
         "5U satellite racks deployable in any room. The 9200L supports IEEE 802.3bt "
         "(PoE++) at up to 60 W per port on selected models, accommodating future "
         "high-powered devices such as PTZ cameras."),

        ("Palo Alto PA-3220 — Hardware Firewall",
         "The assignment mandates a dedicated hardware firewall connected directly to "
         "the border router. The PA-3220 delivers 3.4 Gbps firewall throughput with "
         "full threat prevention enabled, using dedicated processing ASICs that ensure "
         "performance does not degrade under sustained attack traffic (Palo Alto "
         "Networks, 2023). App-ID technology identifies applications by behaviour "
         "rather than port — critical in an intelligence environment where adversaries "
         "may deliberately use non-standard ports or protocol tunnelling to evade "
         "traditional port-based firewall rules. The PA-3220 is a recognised "
         "government-grade security appliance, consistent with the 'security-grade "
         "equipment' requirement."),

        ("Cisco ASR 1001-X — Border Router",
         "The ASR 1001-X supports both a 100 GbE QSFP28 interface (for the DWDM WAN "
         "link) and a 10 GBase-LR SFP+ interface (for the ISP connection) within a "
         "single 1 RU chassis (Cisco Systems, 2021). Hardware IPSec encryption at "
         "20 Gbps ensures inter-site traffic is encrypted at wire speed. BGP routing "
         "enables sophisticated failover policy: the 100 Gbps WDM link to London HQ "
         "is the primary path; the ISP link is the secondary path with appropriate "
         "local preference and MED attributes to control traffic engineering."),
    ]

    for hw_title, hw_body in hw_items:
        add_para(doc, hw_title, bold=True, fontsize=11, colour=(13, 71, 161))
        add_para(doc, hw_body, fontsize=11)

    # ═══════════════════════════════════════════════════════════════════════════
    #  CHAPTER 2 — WAN
    # ═══════════════════════════════════════════════════════════════════════════
    doc.add_page_break()
    add_heading(doc, "Chapter 2: WAN — Optical Network Design", level=1, color=(13, 71, 161))

    # 2.1 Diagram
    add_heading(doc, "2.1 Logical WAN Diagram", level=2)
    add_para(doc,
        "Figure 3 illustrates the logical optical WAN connecting the safe house to "
        "the London main premises. The topology is a 2×3 grid of city nodes, providing "
        "multiple diverse paths and inherent redundancy through the mesh structure."
    )
    add_img(doc, IMG3, width_inches=5.8,
            caption="Figure 3 — Logical WAN Diagram: 100G DWDM Optical Grid with ROADM Nodes and Redundant Paths")

    # 2.2 Topology
    add_heading(doc, "2.2 Topology Analysis", level=2)
    add_para(doc,
        "The provided topology comprises six city nodes arranged in a 2×3 grid: "
        "the safe house at bottom-left and the London main premises (HQ) at top-right, "
        "with four intermediate transit city nodes. This creates two fully disjoint "
        "physical paths between endpoints:"
    )
    add_bullet(doc, "Path A (Primary): Safe House → City E → City D → London HQ")
    add_bullet(doc, "Path B (Backup): Safe House → City C → City A → London HQ")
    add_para(doc,
        "Path disjointness is critical: Path A and Path B share no common city node "
        "or optical fibre span. A complete failure of any single city node or any single "
        "fibre span cannot simultaneously disrupt both paths, ensuring the 100 Gbps "
        "inter-site connection survives any single point of failure (Ramaswami, "
        "Sivarajan and Sasaki, 2010)."
    )

    # 2.3 Design
    add_heading(doc, "2.3 Design Description", level=2)

    add_heading(doc, "WDM Technology", level=3)
    add_para(doc,
        "Dense Wavelength Division Multiplexing (DWDM) multiplexes multiple independent "
        "optical channels onto a single fibre pair, each channel occupying a distinct "
        "wavelength in the C-band (1530–1565 nm), spaced at 100 GHz intervals per "
        "ITU-T G.694.1 (ITU-T, 2012). Each 100 Gbps channel uses Dual Polarisation "
        "Quadrature Phase Shift Keying (DP-QPSK) coherent modulation, which provides "
        "superior tolerance to chromatic dispersion and polarisation mode dispersion "
        "over the 300-mile link distance compared to direct-detect modulation (Saleh "
        "and Teich, 2019)."
    )
    add_para(doc,
        "The WDM-to-router model specified in the assignment brief means that each "
        "endpoint router (safe house and main premises) is equipped with a CFP2-DCO "
        "(Digital Coherent Optics) transceiver installed directly in the router's "
        "line card. The CFP2-DCO converts the router's 100 GbE electrical signal "
        "directly into a coherent DWDM optical wavelength — no separate transponder "
        "chassis is required, reducing equipment footprint and eliminating an additional "
        "device as a potential failure point."
    )

    add_heading(doc, "ROADM Nodes", level=3)
    add_para(doc,
        "Each intermediate city node in the grid hosts a Ciena 6500 "
        "Reconfigurable Optical Add-Drop Multiplexer (ROADM). A ROADM can selectively "
        "add, drop, or pass through individual wavelengths entirely in the optical "
        "domain, without the latency and cost of optical-electrical-optical (OEO) "
        "regeneration (Ciena Corporation, 2023). The Ciena 6500 implements "
        "Colorless, Directionless, Contentionless (CDC) ROADM switching, meaning "
        "any wavelength can be dynamically rerouted to any output port under "
        "software control — this is the mechanism that enables automatic traffic "
        "restoration after a fibre cut."
    )

    add_heading(doc, "Redundancy and Protection Switching", level=3)
    add_para(doc,
        "Under normal operation, Path A carries the live 100 Gbps wavelength "
        "between the safe house and London HQ. Path B is pre-computed and "
        "pre-signalled via GMPLS (Generalised Multi-Protocol Label Switching) "
        "but carries no traffic — this is a 1+1 optical protection scheme. "
        "When the optical layer management system (Ciena MCP) detects a signal "
        "loss event (fibre cut, ROADM failure, or amplifier failure), it "
        "instructs the ROADM nodes to reroute the affected wavelength onto "
        "Path B within 50 milliseconds — faster than TCP retransmission "
        "timers, meaning application sessions at both sites survive the "
        "failure without interruption (ITU-T, 2012)."
    )

    add_heading(doc, "Scalability", level=3)
    add_para(doc,
        "The DWDM architecture scales in two dimensions without physical "
        "infrastructure changes: (1) Additional 100 Gbps wavelengths can be "
        "lit on the existing fibre pairs — the Ciena 6500 supports up to "
        "88 C-band channels, giving a theoretical maximum of 8.8 Tbps per "
        "fibre pair against a current usage of one channel. (2) New city "
        "nodes can be added to the grid by inserting additional ROADM "
        "chassis at the new location and patching it into the existing "
        "fibre plant, extending the mesh without disrupting live traffic "
        "on existing paths."
    )

    # 2.4 Hardware
    add_heading(doc, "2.4 Network Hardware — Critical Evaluation", level=2)

    wan_hw = [
        ("Ciena 6500 Packet-Optical Platform — ROADM Nodes",
         "The Ciena 6500 is a carrier-grade platform deployed in major national and "
         "government optical networks globally (Ciena Corporation, 2023). Its WaveLogic "
         "coherent DSP supports 100G DP-QPSK as a baseline and is software-upgradeable "
         "to 400G DP-16QAM, providing a clear performance evolution path as "
         "inter-site bandwidth demands grow. Integrated Erbium-Doped Fibre Amplifiers "
         "(EDFAs) and Raman pre-amplification extend viable span lengths beyond 120 km "
         "without external amplifier shelves, simplifying the infrastructure required "
         "across the 300-mile route. The MCP network management platform provides "
         "end-to-end wavelength provisioning, real-time performance monitoring, and "
         "automated fault recovery — all critical for an unattended transit node "
         "in a covert location. Alternative platforms such as the ADVA FSP 3000 or "
         "Infinera DTN-X offer similar CDC-ROADM capabilities, but the Ciena 6500's "
         "broader deployment base in UK government and carrier networks provides "
         "better availability of trained field engineers and spare parts."),

        ("Cisco ASR 9000 Series — Endpoint Routers",
         "The ASR 9000's carrier-class modular architecture supports full route "
         "redundancy: the nV Edge clustering technology allows two physical ASR 9000 "
         "chassis to operate as a single logical router, eliminating the router as a "
         "single point of failure at each site endpoint (Cisco Systems, 2022c). "
         "The CFP2-DCO coherent line cards integrate natively with the Ciena 6500 "
         "DWDM layer, and MPLS Traffic Engineering with Fast Reroute (FRR) adds an "
         "IP-layer protection mechanism that complements the 50 ms optical protection "
         "switching — providing dual-layer resilience at both the optical and IP levels. "
         "BGP Route Reflector capability enables sophisticated traffic engineering "
         "policies for optimal use of the primary 100 Gbps WDM path versus the "
         "10 GBase-LR ISP backup link at each site."),
    ]

    for hw_title, hw_body in wan_hw:
        add_para(doc, hw_title, bold=True, fontsize=11, colour=(0, 77, 64))
        add_para(doc, hw_body, fontsize=11)

    # ═══════════════════════════════════════════════════════════════════════════
    #  Conclusion
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "Conclusion", level=1, color=(13, 71, 161))
    add_para(doc,
        "The safe house LAN design implements the three-layer Cisco hierarchical model "
        "with StackWise Virtual core clustering, dual-uplink distribution switching, "
        "and VLAN-segregated access — ensuring that no single component failure can "
        "bring down the network or the physical security systems. The dual-circuit "
        "CCTV architecture ensures continuous surveillance even under targeted attack "
        "against one circuit. The WAN optical design leverages DWDM grid topology and "
        "ROADM-enabled path diversity to deliver a self-healing 100 Gbps inter-site "
        "connection with <50 ms restoration time. Both designs are built entirely from "
        "enterprise and security-grade equipment, are scalable without architectural "
        "redesign, and focus strictly on physical, logical, and architectural "
        "network design as required by the assignment scope."
    )

    # ═══════════════════════════════════════════════════════════════════════════
    #  References
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "References", level=1, color=(13, 71, 161))
    references = [
        "Ciena Corporation (2023) Ciena 6500 Packet-Optical Platform — Product Overview. "
        "Available at: https://www.ciena.com/products/6500 [Accessed 18 April 2026].",

        "Cisco Systems (2021) Cisco ASR 1001-X Router Data Sheet. Available at: "
        "https://www.cisco.com/c/en/us/products/routers/asr-1001-x-router/ "
        "[Accessed 18 April 2026].",

        "Cisco Systems (2022a) Cisco Hierarchical Network Design. Available at: "
        "https://www.cisco.com/c/en/us/td/docs/solutions/Enterprise/Campus/HA_campus_DG/"
        "hacampusdg.html [Accessed 18 April 2026].",

        "Cisco Systems (2022b) Cisco Catalyst 9300 Series Switches Data Sheet. "
        "Available at: https://www.cisco.com/c/en/us/products/collateral/switches/"
        "catalyst-9300-series-switches/nb-06-cat9300-ser-data-sheet-cte-en.html "
        "[Accessed 18 April 2026].",

        "Cisco Systems (2022c) Cisco ASR 9000 Series Aggregation Services Routers. "
        "Available at: https://www.cisco.com/c/en/us/products/routers/asr-9000-series/"
        " [Accessed 18 April 2026].",

        "Cisco Systems (2023a) Cisco Catalyst 9500 Series Switches Data Sheet. "
        "Available at: https://www.cisco.com/c/en/us/products/switches/"
        "catalyst-9500-series-switches/datasheet-listing.html "
        "[Accessed 18 April 2026].",

        "ITU-T (2012) G.694.1: Spectral Grids for WDM Applications — DWDM Frequency "
        "Grid. Geneva: International Telecommunication Union.",

        "Oppenheimer, P. (1999) Top-Down Network Design. Indianapolis: Cisco Press.",

        "Palo Alto Networks (2023) PA-3200 Series Next-Generation Firewalls Datasheet. "
        "Available at: https://www.paloaltonetworks.com/network-security/pa-3200-series "
        "[Accessed 18 April 2026].",

        "Ramaswami, R., Sivarajan, K. and Sasaki, G. (2010) Optical Networks: "
        "A Practical Perspective. 3rd edn. Burlington: Morgan Kaufmann.",

        "Saleh, B.E.A. and Teich, M.C. (2019) Fundamentals of Photonics. "
        "3rd edn. Hoboken: John Wiley & Sons.",
    ]
    for i, ref in enumerate(references, 1):
        rp = doc.add_paragraph(style="List Number")
        rp.add_run(ref).font.size = Pt(10)
        rp.paragraph_format.space_after = Pt(4)

    # ── Save ──────────────────────────────────────────────────────────────────
    out_path = os.path.join(OUT, "SafeHouse_Network_Design_Report.docx")
    doc.save(out_path)
    print(f"[OK] Report saved: {out_path}")
    return out_path


if __name__ == "__main__":
    print("JARVIS — Generating assignment report...")
    p = build_report()
    print(f"\nReport complete: {p}")
